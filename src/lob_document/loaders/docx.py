from __future__ import annotations

import hashlib
import re
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

from lob_document.domain import Block, BlockType, BoundingBox, FigureData, FileIdentity, ImageAsset, Page, SourceDocument, SourceMethod, SourceRef, TableCell, TableData


def _body_items(document: DocumentType):
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _bbox(index: int) -> BoundingBox:
    return BoundingBox(x0=0, y0=index, x1=1, y1=index + 1)


def _table_data(table: Table, index: int) -> tuple[BoundingBox, TableData]:
    rows = len(table.rows)
    columns = max((len(row.cells) for row in table.rows), default=0)
    cells = []
    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            cells.append(TableCell(row=row_index, column=column_index, text="\n".join(p.text for p in cell.paragraphs).strip(), bbox=_bbox(index)))
    return _bbox(index), TableData(row_count=max(1, rows), column_count=max(1, columns), cells=cells)


def _add_block(blocks: list[Block], document_id: str, index: int, block_type: BlockType, text: str, heading_level: int | None = None, table: TableData | None = None) -> None:
    text = text.strip()
    if not text and table is None:
        return
    bbox = _bbox(index)
    digest = hashlib.sha256(f"{index}:{block_type}:{text}".encode()).hexdigest()[:16]
    ref = SourceRef(document_id=document_id, page_number=1, bbox=bbox, line_start=index + 1, line_end=index + 1)
    blocks.append(Block(id=f"{document_id}_block_{digest}", type=block_type, text=text or None, bbox=bbox, reading_order=len(blocks), heading_level=heading_level, source_method=SourceMethod.NATIVE, source_engine="python-docx", source_engine_version="1", source_ref=ref, table=table))


def load_docx(path: Path, artifacts_dir: Path | None = None) -> SourceDocument:
    source_path = path.expanduser().resolve()
    if not source_path.is_file():
        raise FileNotFoundError(f"file does not exist: {source_path}")
    if source_path.suffix.lower() != ".docx":
        raise ValueError("only DOCX input is supported by the Word loader")
    data = source_path.read_bytes()
    digest = hashlib.sha256(data).hexdigest()
    document_id = f"doc_{digest[:24]}"
    document = Document(source_path)
    blocks: list[Block] = []
    for index, item in enumerate(_body_items(document)):
        if isinstance(item, Paragraph):
            style = item.style.name if item.style else ""
            match = re.search(r"(?:Heading|标题)\s*([1-6])", style, re.IGNORECASE)
            block_type = BlockType.HEADING if match else BlockType.LIST if item.style and "List" in item.style.name else BlockType.PARAGRAPH
            _add_block(blocks, document_id, index, block_type, item.text, int(match.group(1)) if match else None)
        else:
            bbox, table = _table_data(item, index)
            _add_block(blocks, document_id, index, BlockType.TABLE, "", table=table)

    if artifacts_dir is not None:
        assets_dir = artifacts_dir / "assets"
        assets_dir.mkdir(parents=True, exist_ok=True)
    else:
        assets_dir = None
    for shape in document.inline_shapes:
        blip = shape._inline.graphic.graphicData.pic.blipFill.blip
        related = document.part.related_parts.get(blip.embed)
        if related is None or not hasattr(related, "blob"):
            continue
        image_data = related.blob
        image_digest = hashlib.sha256(image_data).hexdigest()
        extension = related.content_type.split("/")[-1].replace("jpeg", "jpg")
        filename = f"asset_{image_digest[:24]}.{extension}"
        asset_path = Path("assets") / filename
        if assets_dir is not None:
            (assets_dir / filename).write_bytes(image_data)
        else:
            asset_path = source_path.parent / filename
        index = len(blocks)
        bbox = _bbox(index)
        ref = SourceRef(document_id=document_id, page_number=1, bbox=bbox, line_start=index + 1, line_end=index + 1)
        blocks.append(Block(id=f"{document_id}_figure_{image_digest[:16]}", type=BlockType.FIGURE, bbox=bbox, reading_order=index, source_method=SourceMethod.NATIVE, source_engine="python-docx", source_engine_version="1", source_ref=ref, figure=FigureData(asset=ImageAsset(id=f"asset_{image_digest[:24]}", path=str(asset_path), sha256=image_digest, media_type=related.content_type, width=1, height=1))))

    page = Page(id=f"{document_id}_page_0001", page_number=1, width=1, height=max(1, len(blocks)), blocks=blocks)
    return SourceDocument(source=FileIdentity(id=document_id, sha256=digest, filename=source_path.name, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", size_bytes=len(data)), page_count=1, pages=[page])
